from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.pdfgen import canvas
import textwrap


class CampaignExporter:
    """Export cold email campaigns to professional DOCX and PDF formats"""
    
    def __init__(self):
        self.purple = RGBColor(124, 58, 237)  # #7C3AED
        self.dark_purple = RGBColor(91, 33, 182)  # #5B21B6
        self.green = RGBColor(16, 185, 129)  # #10B981
        self.red = RGBColor(239, 68, 68)  # #EF4444
        self.gray = RGBColor(107, 114, 128)  # #6B7280
    
    def export_to_docx(self, campaign_data: dict) -> BytesIO:
        """Export campaign to beautiful DOCX format"""
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Cold Email Campaign Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = self.purple
        title_run.font.size = Pt(28)
        title_run.bold = True
        
        doc.add_paragraph()
        
        # Company Info
        company = campaign_data.get('company', {})
        info = doc.add_paragraph()
        info.add_run('Target Company: ').bold = True
        info.add_run(f"{company.get('name', 'N/A')}\n")
        info.add_run('Industry: ').bold = True
        info.add_run(f"{company.get('industry', 'N/A')}\n")
        info.add_run('Company Size: ').bold = True
        info.add_run(f"{company.get('size', 'N/A')}\n")
        info.add_run('Generated: ').bold = True
        info.add_run(datetime.utcnow().strftime('%B %d, %Y at %H:%M UTC'))
        
        doc.add_paragraph()
        
        # Strategic Analysis
        doc.add_heading('📊 Strategic Analysis', level=1)
        analysis = campaign_data.get('analysis', {}).get('strategic_brief', {})
        
        # Pain Points
        doc.add_heading('Top 3 Pain Points', level=2)
        for i, pain in enumerate(analysis.get('top_3_pain_points', []), 1):
            pain_text = self._to_string(pain)
            p = doc.add_paragraph(f"{i}. {pain_text}", style='List Bullet')
        
        # Objections
        doc.add_heading('Key Objections', level=2)
        for obj in analysis.get('key_objections', []):
            obj_text = self._to_string(obj)
            doc.add_paragraph(f"• {obj_text}", style='List Bullet')
        
        # Value Props
        doc.add_heading('Resonant Value Propositions', level=2)
        for vp in analysis.get('resonant_value_propositions', []):
            vp_text = self._to_string(vp)
            doc.add_paragraph(f"• {vp_text}", style='List Bullet')
        
        doc.add_page_break()
        
        # Cold Emails
        doc.add_heading('📧 Cold Email Campaigns', level=1)
        
        emails = campaign_data.get('cold_emails', {})
        for approach, email_data in emails.items():
            # Approach title
            approach_title = approach.replace('_', ' ').title()
            heading = doc.add_heading(f"Approach: {approach_title}", level=2)
            heading.runs[0].font.color.rgb = self.purple
            
            # Subject
            subject_para = doc.add_paragraph()
            subject_para.add_run('SUBJECT: ').bold = True
            subject_para.add_run(email_data.get('subject', ''))
            
            # Email body
            email_body = email_data.get('email', '')
            doc.add_paragraph(email_body)
            
            # Variants
            variants = email_data.get('subject_variants', [])
            if variants:
                var_para = doc.add_paragraph()
                var_para.add_run('Alternative Subjects:\n').bold = True
                for i, variant in enumerate(variants, 1):
                    var_para.add_run(f"  {i}. {variant}\n")
            
            doc.add_paragraph()
        
        doc.add_page_break()
        
        # Follow-ups
        doc.add_heading('📬 Follow-Up Sequence', level=1)
        
        followups = campaign_data.get('followup_sequence', [])
        for followup in followups:
            day = followup.get('day', 0)
            heading = doc.add_heading(f"Follow-up {followups.index(followup) + 1} (Day {day})", level=2)
            heading.runs[0].font.color.rgb = self.green
            
            subject_para = doc.add_paragraph()
            subject_para.add_run('SUBJECT: ').bold = True
            subject_para.add_run(followup.get('subject', ''))
            
            doc.add_paragraph(followup.get('body', ''))
            doc.add_paragraph()
        
        doc.add_page_break()
        
        # Recommendations
        doc.add_heading('💡 Strategic Recommendations', level=1)
        recommendations = campaign_data.get('recommendations', {}).get('strategic_recommendations', '')
        doc.add_paragraph(recommendations)
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def export_to_pdf(self, campaign_data: dict) -> BytesIO:
        """Export campaign to AWARD-WINNING PDF format"""
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=50,
            leftMargin=50,
            topMargin=50,
            bottomMargin=50
        )
        
        # Container for elements
        elements = []
        
        # Styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=28,
            textColor=colors.HexColor('#7C3AED'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        section_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=18,
            textColor=colors.HexColor('#5B21B6'),
            spaceAfter=12,
            spaceBefore=20,
            fontName='Helvetica-Bold'
        )
        
        subsection_style = ParagraphStyle(
            'SubSection',
            parent=styles['Heading3'],
            fontSize=14,
            textColor=colors.HexColor('#7C3AED'),
            spaceAfter=8,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            leading=16,
            textColor=colors.black,
            alignment=TA_JUSTIFY,
            fontName='Helvetica'
        )
        
        email_style = ParagraphStyle(
            'EmailBody',
            parent=styles['BodyText'],
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#1F2937'),
            alignment=TA_LEFT,
            fontName='Helvetica',
            leftIndent=10,
            rightIndent=10
        )
        
        # Title
        elements.append(Paragraph("Cold Email Campaign Report", title_style))
        elements.append(Spacer(1, 0.3*inch))
        
        # Company Info
        company = campaign_data.get('company', {})
        info_text = f"""
        <b>Target Company:</b> {self._escape(company.get('name', 'N/A'))}<br/>
        <b>Industry:</b> {self._escape(company.get('industry', 'N/A'))}<br/>
        <b>Company Size:</b> {self._escape(company.get('size', 'N/A'))}<br/>
        <b>Generated:</b> {datetime.utcnow().strftime('%B %d, %Y at %H:%M UTC')}
        """
        elements.append(Paragraph(info_text, body_style))
        elements.append(Spacer(1, 0.4*inch))
        
        # Strategic Analysis
        elements.append(Paragraph("📊 Strategic Analysis", section_style))
        analysis = campaign_data.get('analysis', {}).get('strategic_brief', {})
        
        # Pain Points
        elements.append(Paragraph("■ Top 3 Pain Points", subsection_style))
        for i, pain in enumerate(analysis.get('top_3_pain_points', []), 1):
            pain_text = self._to_string(pain)
            wrapped = self._wrap_text(pain_text, 90)
            elements.append(Paragraph(f"{i}. {self._escape(wrapped)}", body_style))
            elements.append(Spacer(1, 0.1*inch))
        
        # Objections
        elements.append(Spacer(1, 0.2*inch))
        elements.append(Paragraph("■ Key Objections", subsection_style))
        for obj in analysis.get('key_objections', []):
            obj_text = self._to_string(obj)
            wrapped = self._wrap_text(obj_text, 90)
            elements.append(Paragraph(f"• {self._escape(wrapped)}", body_style))
            elements.append(Spacer(1, 0.08*inch))
        
        # Value Props
        elements.append(Spacer(1, 0.2*inch))
        elements.append(Paragraph("■ Resonant Value Propositions", subsection_style))
        for vp in analysis.get('resonant_value_propositions', []):
            vp_text = self._to_string(vp)
            wrapped = self._wrap_text(vp_text, 90)
            elements.append(Paragraph(f"• {self._escape(wrapped)}", body_style))
            elements.append(Spacer(1, 0.08*inch))

        # Spacer before next section instead of hard page break
        elements.append(Spacer(1, 0.5*inch))
        
        # Cold Emails
        elements.append(Paragraph("📧 Cold Email Campaigns", section_style))
        elements.append(Spacer(1, 0.2*inch))
        
        emails = campaign_data.get('cold_emails', {})
        for approach, email_data in emails.items():
            approach_title = approach.replace('_', ' ').title()
            elements.append(Paragraph(f"Approach: {approach_title}", subsection_style))
            
            # Subject
            subject = self._escape(email_data.get('subject', ''))
            elements.append(Paragraph(f"<b>SUBJECT:</b> {subject}", body_style))
            elements.append(Spacer(1, 0.1*inch))
            
            # Email body with wrapping
            email_body = self._to_string(email_data.get('email', ''))
            wrapped_body = self._wrap_text(email_body, 85)
            
            # Create table for email body (better formatting)
            email_table = Table(
                [[Paragraph(self._escape(wrapped_body), email_style)]],
                colWidths=[6.5*inch]
            )
            email_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F9FAFB')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 12),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('LEFTPADDING', (0, 0), (-1, -1), 12),
                ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#E5E7EB')),
            ]))
            elements.append(email_table)
            
            # Variants
            variants = email_data.get('subject_variants', [])
            if variants:
                elements.append(Spacer(1, 0.1*inch))
                elements.append(Paragraph("<b>Alternative Subjects:</b>", body_style))
                for i, variant in enumerate(variants, 1):
                    variant_text = self._to_string(variant)
                    wrapped_var = self._wrap_text(variant_text, 85)
                    elements.append(Paragraph(f"  {i}. {self._escape(wrapped_var)}", body_style))

            elements.append(Spacer(1, 0.4*inch))

        # Follow-ups
        elements.append(Paragraph("📬 Follow-Up Sequence", section_style))
        elements.append(Spacer(1, 0.2*inch))
        
        followups = campaign_data.get('followup_sequence', [])
        for i, followup in enumerate(followups, 1):
            day = followup.get('day', 0)
            elements.append(Paragraph(f"Follow-up {i} (Day {day})", subsection_style))
            
            # Subject
            subject = self._escape(self._to_string(followup.get('subject', '')))
            elements.append(Paragraph(f"<b>SUBJECT:</b> {subject}", body_style))
            elements.append(Spacer(1, 0.1*inch))
            
            # Body
            body = self._to_string(followup.get('body', ''))
            wrapped_body = self._wrap_text(body, 85)
            
            followup_table = Table(
                [[Paragraph(self._escape(wrapped_body), email_style)]],
                colWidths=[6.5*inch]
            )
            followup_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#ECFDF5')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 12),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('LEFTPADDING', (0, 0), (-1, -1), 12),
                ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#10B981')),
            ]))
            elements.append(followup_table)
            elements.append(Spacer(1, 0.4*inch))
        
        # Recommendations
        elements.append(Paragraph("💡 Strategic Recommendations", section_style))
        elements.append(Spacer(1, 0.2*inch))

        recommendations = self._to_string(campaign_data.get('recommendations', {}).get('strategic_recommendations', ''))

        # Split recommendations into paragraphs for better readability
        rec_paragraphs = recommendations.split('\n\n')
        for rec_para in rec_paragraphs:
            if rec_para.strip():
                wrapped_rec = self._wrap_text(rec_para.strip(), 95)
                elements.append(Paragraph(self._escape(wrapped_rec), body_style))
                elements.append(Spacer(1, 0.15*inch))  # Add spacing between paragraphs
        
        # Footer
        elements.append(Spacer(1, 0.5*inch))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#6B7280'),
            alignment=TA_CENTER
        )
        elements.append(Paragraph(
            f"Generated by Cold Email Generator Pro | {datetime.utcnow().strftime('%B %d, %Y')}",
            footer_style
        ))
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        
        return buffer
    
    def _to_string(self, value) -> str:
        """Convert any value to string safely, handling dicts properly"""
        if value is None:
            return ""
        if isinstance(value, str):
            return value
        if isinstance(value, dict):
            # Extract meaningful content from dict objects
            parts = []

            # Common fields to extract in priority order
            priority_fields = ['pain_point', 'objection', 'value_proposition', 'description', 'why_it_works', 'urgency', 'text', 'content']

            # First add priority fields if they exist
            for field in priority_fields:
                if field in value and value[field]:
                    parts.append(str(value[field]))

            # If no priority fields found, add all non-empty values
            if not parts:
                for k, v in value.items():
                    if v and not k.startswith('_'):  # Skip empty and private fields
                        parts.append(f"{k.replace('_', ' ').title()}: {v}")

            return ' - '.join(parts) if parts else str(value)

        if isinstance(value, (list, tuple)):
            # Handle lists/tuples
            return ', '.join(self._to_string(item) for item in value)

        return str(value)
    
    def _wrap_text(self, text: str, width: int = 90) -> str:
        """Wrap text to prevent overflow"""
        if not text:
            return ""
        
        # Ensure it's a string
        text = self._to_string(text)
        
        lines = text.split('\n')
        wrapped_lines = []
        
        for line in lines:
            if len(line) <= width:
                wrapped_lines.append(line)
            else:
                wrapped = textwrap.fill(line, width=width, break_long_words=False, break_on_hyphens=False)
                wrapped_lines.append(wrapped)
        
        return '\n'.join(wrapped_lines)
    
    def _escape(self, text: str) -> str:
        """Escape XML characters for reportlab"""
        if not text:
            return ""
        
        # Ensure it's a string
        text = self._to_string(text)
        
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        text = text.replace('"', '&quot;')
        text = text.replace("'", '&apos;')
        
        return text